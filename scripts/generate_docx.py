import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('InfraMinds: Technical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph()
    p.add_run('InfraMinds').bold = True
    p.add_run(' is an autonomous infrastructure agent that transforms natural language intent (or images) into a verified "Living State Graph" of cloud architecture. Unlike standard text-to-code tools, it employs a ')
    p.add_run('multi-phase reasoning engine').bold = True
    p.add_run(' (Intent -> Reasoned -> Implementation) to visualize, simulate, and verify infrastructure changes before generating Terraform code. It features a "Glass Box" approach, allowing users to see the "Blast Radius" of changes and relying on autonomous self-healing loops to ensure correctness.')

    # Core Architecture
    doc.add_heading('2. Core Architecture & Tech Stack', level=1)
    
    doc.add_heading('2.1 Backend (/backend)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Framework: ').bold = True
    p.add_run('FastAPI (Python)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AI Engine: ').bold = True
    p.add_run('Google Gemini (via google-genai SDK), utilizing models like gemini-2.0-flash.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Graph Engine: ').bold = True
    p.add_run('NetworkX for maintaining the dependency graph and performing topological analysis (Blast Radius).')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Infrastructure Engine: ').bold = True
    p.add_run('Terraform (via tflocal for LocalStack simulation).')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Pipeline: ').bold = True
    p.add_run('A robust 5-stage pipeline for safe execution.')

    doc.add_heading('2.2 Frontend (/frontend)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Framework: ').bold = True
    p.add_run('Next.js (React 19)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Visualization: ').bold = True
    p.add_run('React Flow with Dagre for auto-layout.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('State Management: ').bold = True
    p.add_run('Real-time streaming of agent thoughts and graph updates.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Styling: ').bold = True
    p.add_run('Tailwind CSS + Framer Motion.')

    # The 4-Phase Generation Flow
    doc.add_heading('3. The 4-Phase Generation Flow', level=1)
    doc.add_paragraph('InfraMinds does not jump straight to code. It evolves the architecture through 4 distinct phases to ensure semantically correct and improved infrastructure.')

    doc.add_heading('Phase 1: Intent Generation (Abstract)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Input: ').bold = True
    p.add_run('Natural Language (e.g., "I need a scalable e-commerce app") or Architecture Diagrams (Images).')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Process: ').bold = True
    p.add_run('\n')
    p.add_run('   - Text: Uses get_intent_text_prompt to map request to Semantic Types only (e.g., compute_service, relational_database).\n')
    p.add_run('   - Image: Uses get_vision_prompt to extract architecture from diagrams.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Output: ').bold = True
    p.add_run('Intent Graph (Method: generate_intent_stream).')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Goal: ').bold = True
    p.add_run('Capture *what* the user wants without getting bogged down in *how* (AWS primitives).')

    doc.add_heading('Phase 2: Reasoning & Policy Enforcement', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Input: ').bold = True
    p.add_run('Intent Graph')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Process: ').bold = True
    p.add_run('\n')
    p.add_run('   - The Policy Engine (apply_policies_gen) acts as a "Cloud Architect".\n')
    p.add_run('   - Enforces baseline policies:\n')
    p.add_run('       • Isolation: DBs/Caches must not be public.\n')
    p.add_run('       • Least Privilege: Strict connectivity.\n')
    p.add_run('       • Ingress Discipline: Public compute must use Load Balancers.\n')
    p.add_run('   - Self-Healing Loop: If violations are found, the graph is mutated (e.g., removing a direct public edge to a DB) and re-evaluated (up to 3 cycles).')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Output: ').bold = True
    p.add_run('Reasoned Graph (still semantic, but policy-compliant).')

    doc.add_heading('Phase 3: Cloud Expansion (Implementation)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Input: ').bold = True
    p.add_run('Reasoned Graph')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Process: ').bold = True
    p.add_run('\n')
    p.add_run('   - The Platform Engineer agent (expand_architecture_gen) maps semantic nodes to concrete AWS Primitives.\n')
    p.add_run('   - Examples: compute_service -> aws_instance, relational_database -> aws_db_instance.\n')
    p.add_run('   - Materialization: Adds necessary supporting infrastructure (VPCs, Subnets, Internet Gateways, Route Tables) that wasn\'t explicitly asked for but is required.\n')
    p.add_run('   - Preservation Check: Ensures all semantic nodes from Phase 2 exist in Phase 3.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Output: ').bold = True
    p.add_run('Implementation Graph (The "Living State Graph").')

    doc.add_heading('Phase 4: Optimization & Cost', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Input: ').bold = True
    p.add_run('Implementation Graph')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Process: ').bold = True
    p.add_run('\n')
    p.add_run('   - Cost Estimator: CostEstimator (in cost.py) analyzes resource types (instance_type, volume_size).\n')
    p.add_run('   - Estimates monthly cost using AI knowledge of on-demand pricing.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Output: ').bold = True
    p.add_run('Cost Report JSON (displayed in UI).')

    # Self-Healing Loops
    doc.add_heading('4. Self-Healing Loops', level=1)
    doc.add_paragraph('The system is designed to be autonomous and resilient through multiple feedback loops:')
    
    doc.add_heading('1. Graph Policy Loop (Phase 2)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Trigger: ').bold = True
    p.add_run('Policy violation (e.g., "Database exposed to internet").')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Action: ').bold = True
    p.add_run('LLM modifies the graph structure (removes edge, adds attributes).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Retry: ').bold = True
    p.add_run('Re-runs policy check until pure.')

    doc.add_heading('2. Architecture Expansion Loop (Phase 3)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Trigger: ').bold = True
    p.add_run('Monotonicity failure (dropped nodes) or Hallucination (abstract nodes remaining).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Action: ').bold = True
    p.add_run('Retries expansion prompt.')

    doc.add_heading('3. Terraform Pipeline Loop (Execution)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Trigger: ').bold = True
    p.add_run('terraform validate or terraform apply failure.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Action: ').bold = True
    p.add_run('\n')
    p.add_run('   - Captures stderr.\n')
    p.add_run('   - Calls _fix_code in pipeline.py.\n')
    p.add_run('   - LLM analyzes the error ("X-ray vision"), proposes a fix, and generates patched HCL.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Retry: ').bold = True
    p.add_run('Re-runs the failed stage (up to 3 times).')

    # Knowledge Graph
    doc.add_heading('5. Usage of Knowledge Graph', level=1)
    doc.add_paragraph('The project relies heavily on NetworkX (agent.py) to maintain a "Living State":')
    
    p = doc.add_paragraph(style='List Number')
    p.add_run('Dependency Tracking: ').bold = True
    p.add_run('Knows that Subnet A is inside VPC B, and Instance C is inside Subnet A.')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Blast Radius Analysis: ').bold = True
    p.add_run('\n')
    p.add_run('   - Uses graph traversal (descendants/ancestors) to simulate failure impact.\n')
    p.add_run('   - Prompts (get_blast_radius_prompt) combine structural graph data with AI reasoning (e.g., "Deleting this Security Group allows 0.0.0.0/0").')

    p = doc.add_paragraph(style='List Number')
    p.add_run('Visual Layout: ').bold = True
    p.add_run('The graph structure drives the React Flow visualization, grouping resources by Subnet/VPC (using dagre for auto-layout).')

    # The 5-Stage Testing Plan
    doc.add_heading('6. The 5-Stage Testing Plan (Pipeline)', level=1)
    doc.add_paragraph('To ensure the correctness of the generated Terraform code, pipeline.py executes a strict 5-stage plan:')

    doc.add_heading('Stage 1 (Setup/Clean)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Removes .terraform.lock.hcl, state files, and provider overrides to ensure a clean slate.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Runs terraform init.')

    doc.add_heading('Stage 2 (Validate)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Runs terraform validate.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Performs Static Policy Check (_check_policy): Scans HCL/Regex for forbidden patterns (e.g., inline ingress rules in SG).')

    doc.add_heading('Stage 3 (Plan)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Runs tflocal plan.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Verifies that the code is logically consistent.')

    doc.add_heading('Stage 4 (Apply)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Runs tflocal apply -auto-approve.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Actually creates resources in the LocalStack simulation.')

    doc.add_heading('Stage 5 (Verify)', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Runs a Python test script (test_infra.py).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Logic: The script (generated by AI) performs "Black Box" testing:\n')
    p.add_run('   - Checks if resources exist (via Boto3).\n')
    p.add_run('   - Checks connectivity (HTTP 200).\n')
    p.add_run('   - Validates configuration (e.g., "Is port 80 open?").')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('If this script fails, the pipeline is marked as failed even if Terraform applied successfully.')

    # Blast Radius & Time Travel
    doc.add_heading('7. Blast Radius & Time Travel', level=1)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Concept: ').bold = True
    p.add_run('Interactive impact analysis before deployment.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Mechanism: ').bold = True
    p.add_run('\n')
    p.add_run('   1. User selects a node (e.g., "NAT Gateway").\n')
    p.add_run('   2. Agent traverses the NetworkX graph to find all dependent downstream nodes.\n')
    p.add_run('   3. AI Agent adds semantic reasoning (e.g., "Removing NAT Gateway breaks internet access for Private Subnets").')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Visualization: ').bold = True
    p.add_run('\n')
    p.add_run('   - Target Node: Pulsing Red.\n')
    p.add_run('   - Affected Nodes: Orange Glow with dashed orange edges showing the "Kill Chain".')

    # Cost Estimation
    doc.add_heading('8. Cost Estimation', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Logic: ').bold = True
    p.add_run('backend/cost.py.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Method: ').bold = True
    p.add_run('Extracts resource properties (t2.micro, gp3) -> Sends to LLM ("You are an AWS Pricing Expert") -> Returns JSON breakdown.')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Display: ').bold = True
    p.add_run('Tablet in the UI showing Total Monthly Cost and itemized breakdown.')

    # Footer
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Generated by InfraMinds Autonomous Agent')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save('/Users/yashrajput/Desktop/InfraMinds/doc/project_documentation.docx')
    print("Document created successfully at /Users/yashrajput/Desktop/InfraMinds/doc/project_documentation.docx")

if __name__ == "__main__":
    create_document()
